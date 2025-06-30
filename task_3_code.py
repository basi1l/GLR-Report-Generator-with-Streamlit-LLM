import streamlit as st
import fitz
import docx
import requests
import re
from io import BytesIO

st.set_page_config(page_title="GLR Generator - USAA / Wayne / GuideOne", layout="wide")
st.title("📝 GLR Report Generator – USAA / Wayne / GuideOne")

OPENROUTER_API_KEY = st.secrets.get("OPENROUTER_API_KEY", "")

template_file = st.file_uploader(
    "📄 Upload Template (.docx)",
    type=["docx"],
    help="Upload the empty GLR template DOCX file for USAA, Wayne or GuideOne."
)

photo_report = st.file_uploader(
    "📷 Upload Photo Report (.pdf)",
    type=["pdf"],
    help="Upload the inspection report (PDF) that describes the damage."
)

def extract_pdf_text(pdf_file):
    text = ""
    with fitz.open(stream=pdf_file.read(), filetype="pdf") as doc:
        for page in doc:
            text += page.get_text()
    return text

def detect_template_type(docx_stream):
    doc = docx.Document(docx_stream)
    full_text = []
    for p in doc.paragraphs:
        full_text.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)
    combined = "\n".join(full_text)
    if "[XM8_" in combined and "GUIDEONE" in combined.upper():
        return "guideone"
    elif "[XM8_" in combined:
        return "wayne"
    elif "[INSURED_NAME]" in combined:
        return "usaa"
    else:
        return "unknown"

def build_prompt_usaa(pdf_text):
    fields = [
        "INSURED_NAME", "DATE_LOSS", "DATE_RECEIVED", "DATE_INSPECTED",
        "MORTGAGEE", "INSURED_H_STREET", "INSURED_H_CITY", "INSURED_H_STATE",
        "INSURED_H_ZIP", "DWELLING_DESCRIPTION", "PROPERTY_CONDITION", "INSPECTION_SUMMARY",
        "DWELLING_SECTION", "ELEVATION_SECTION", "INTERIOR_SECTION", "OTHER_STRUCTURES_SECTION",
        "CONTENTS_SECTION", "REVIEW_SECTION", "SUPPLEMENT_SECTION", "PRIORS_SECTION",
        "CODE_ITEMS_SECTION", "OVERHEAD_PROFIT_SECTION", "MICA_QA_SECTION", "MORTGAGEE_SECTION",
        "CAUSE_ORIGIN_SECTION", "SUBROGATION_SECTION", "SALVAGE_SECTION"
    ]
    prompt = "You are a professional insurance adjuster assistant. Based on the photo report below, extract each of the following fields and return them in the format:\n\n"
    for f in fields:
        prompt += f"[{f}]:\n"
    prompt += "\nIf any field is not available, respond with 'Not observed'.\n\nPHOTO REPORT:\n" + pdf_text
    return prompt

def build_prompt_wayne(pdf_text):
    fields = [
        "XM8_INSURED_NAME", "XM8_DATE_LOSS", "XM8_DATE_INSPECTED", "XM8_ESTIMATOR_NAME",
        "XM8_ESTIMATOR_E_MAIL", "XM8_ESTIMATOR_C_PHONE", "XM8_CLAIM_NUMBER", "XM8_LOCATION_LOSS",
        "XM8_RISK_INFO", "XM8_CAUSE_ORIGIN", "XM8_INSPECTION_SUMMARY", "XM8_INTERIOR_SECTION",
        "XM8_CONTENTS_SECTION", "XM8_REVIEW_SECTION", "XM8_SUPPLEMENT_SECTION", "XM8_SALVAGE_SECTION",
        "XM8_SUBROGATION_SECTION", "XM8_DATE_CURRENT"
    ]
    prompt = "You are a claims adjuster assistant. Based on the photo report below, return each of the following Wayne GLR fields in the format:\n\n"
    for f in fields:
        prompt += f"[{f}]:\n"
    prompt += "\nIf any field is not found, return 'Not observed'.\n\nPHOTO REPORT:\n" + pdf_text
    return prompt

def build_prompt_guideone(pdf_text):
    fields = [
        "XM8_INSURED_NAME", "XM8_DATE_LOSS", "XM8_DATE_INSPECTED", "XM8_DATE_CURRENT",
        "XM8_ESTIMATOR_NAME", "XM8_ESTIMATOR_E_MAIL", "XM8_ESTIMATOR_C_PHONE", "XM8_INSURED_P_STREET",
        "XM8_INSURED_P_CITY", "XM8_INSURED_P_STATE", "XM8_INSURED_P_ZIP", "XM8_CLAIM_NUMBER",
        "XM8_LOCATION_LOSS", "XM8_RISK_INFO", "XM8_CAUSE_ORIGIN", "XM8_INTERIOR_SECTION",
        "XM8_SALVAGE_SECTION", "XM8_SUBROGATION_SECTION", "XM8_REVIEW_SECTION"
    ]
    prompt = "You are an expert adjuster assistant. Based on the report below, return each of the following GuideOne GLR fields in the format:\n\n"
    for f in fields:
        prompt += f"[{f}]:\n"
    prompt += "\nIf a field is not present in the text, return 'Not observed'.\n\nPHOTO REPORT:\n" + pdf_text
    return prompt

def query_llm(prompt):
    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "mistralai/mistral-7b-instruct",
        "messages": [{"role": "user", "content": prompt}]
    }
    res = requests.post("https://openrouter.ai/api/v1/chat/completions", json=payload, headers=headers)
    if res.status_code != 200:
        st.error(f"❌ API Error: {res.status_code}")
        st.code(res.text)
        return "ERROR"
    try:
        return res.json()["choices"][0]["message"]["content"]
    except Exception as e:
        st.error("❌ Failed to parse LLM response.")
        st.code(res.text)
        raise e

def parse_llm_response(output):
    pattern = r"(?m)^\[([A-Z0-9_]+)\]:\s*\n?(.+?)(?=\n\[|$)"
    return {key.strip(): val.strip() for key, val in re.findall(pattern, output, re.DOTALL)}

def fill_docx_template(template_stream, fields):
    doc = docx.Document(template_stream)
    for para in doc.paragraphs:
        for key, val in fields.items():
            token = f"[{key}]"
            if token in para.text:
                for run in para.runs:
                    if token in run.text:
                        run.text = run.text.replace(token, val)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in fields.items():
                    token = f"[{key}]"
                    if token in cell.text:
                        cell.text = cell.text.replace(token, val)
    return doc

# Main logic
if template_file and photo_report:
    if not template_file.name.lower().endswith(".docx"):
        st.error("❌ Please upload a valid .docx template file.")
        st.stop()

    with st.spinner("📖 Reading report..."):
        pdf_text = extract_pdf_text(photo_report)

    with st.spinner("🔍 Detecting template type..."):
        template_type = detect_template_type(template_file)

    if template_type == "unknown":
        st.error("❌ Template format not recognized.")
    else:
        st.success(f"✅ Detected: {template_type.upper()} template")
        if template_type == "usaa":
            prompt = build_prompt_usaa(pdf_text)
        elif template_type == "wayne":
            prompt = build_prompt_wayne(pdf_text)
        else:
            prompt = build_prompt_guideone(pdf_text)

        with st.spinner("🤖 Calling LLM..."):
            llm_response = query_llm(prompt)

        if llm_response != "ERROR":
            st.expander("📋 LLM Output").write(llm_response)
            parsed = parse_llm_response(llm_response)

            with st.spinner("📝 Filling template..."):
                filled_doc = fill_docx_template(template_file, parsed)
                buf = BytesIO()
                filled_doc.save(buf)
                buf.seek(0)
                st.download_button("📥 Download Completed Report", buf, file_name="generated_glr.docx")
